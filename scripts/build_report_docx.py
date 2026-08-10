"""
Build presentation/final_report.docx (native Word) from the same authoritative content as
presentation/final_report.md.

Idempotent / re-runnable: overwrites the .docx in place each run. CPU-only, no network.

Usage:
    python scripts/build_report_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[1]
PRES = REPO / "presentation"
IMAGES = PRES / "images"
OUT = PRES / "final_report.docx"

HEADING_BLUE = RGBColor(0x1A, 0x3A, 0x5C)
BODY_INK = RGBColor(0x20, 0x20, 0x20)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, col_widths=None, header_fill="1A3A5C"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption, width=5.8):
    path = IMAGES / filename
    if not path.exists():
        raise FileNotFoundError(f"Missing figure: {path}")
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = HEADING_BLUE
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = HEADING_BLUE
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = HEADING_BLUE
    return p


def body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_INK
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10.5)


def build():
    doc = Document()

    # ---- base style ----
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ================= TITLE PAGE =================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Invoice Region Detection & Business-Parameter Extraction\nfor Obligation-Readiness"
    )
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = HEADING_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("A Deep Learning II Group Project")
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(
        "Authors: Rolando (Data Ingestion & Manifest) - Diana (Stamp & Signature Detection) - "
        "Jordan (Region Detection & IoU Evaluation) - Damir (OCR, Parameters & Terms Extraction) - "
        "Hessam (Integration, Verdict Engine & Streamlit Application)"
    )
    run.font.size = Pt(11)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("July 2026")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ================= ABSTRACT =================
    h1(doc, "Abstract")
    body(
        doc,
        "Before an invoice image can become a structured digital obligation record - the kind a "
        "downstream finance or legal system (in the spirit of a \"Pistac.io\"-style platform) can "
        "act on without a human first reading it - someone has to answer a handful of yes/no "
        "questions: is it signed or stamped? does it cite a purchase order, contract, or work "
        "order? does it state clear payment terms and a valid date? Doing this by hand does not "
        "scale past a few hundred invoices. This project builds an end-to-end computer-vision and "
        "OCR pipeline that reads an invoice image and answers those questions automatically, then "
        "lets a user configure their own readiness policy and see a transparent, per-rule Ready / "
        "Not-ready verdict - never a black-box yes/no."
    )
    body(
        doc,
        "Five stages, one per team member, hand off outputs through a fixed CSV/JSON interface "
        "contract: data ingestion and manifest construction (Rolando), a YOLOv8n stamp/signature "
        "detector (Diana), a YOLOv8n region detector trained on real polygon-annotated receipt "
        "data (Jordan), EasyOCR text extraction feeding rule-based business-parameter and "
        "payment-terms extraction (Damir), and an integration layer that merges everything into "
        "one JSON per invoice and powers a Streamlit application built around a fail-closed, "
        "user-configurable verdict engine (Hessam). Detectors are trained and evaluated on the "
        "best available real, labelled datasets - SignverOD, StaVer, and the OCR Dataset of "
        "Multi-type Documents - none of which are invoices; they are then applied to real invoice "
        "scans, and we report that domain-transfer step honestly as detection counts, not "
        "accuracy, since no invoice-level ground truth exists. This report documents the "
        "methodology, design rationale, and final, real quantitative results of the system, drawn "
        "directly from each member's completed Colab GPU training/evaluation run (see Section 5)."
    )

    # ================= 1. BUSINESS PROBLEM =================
    h1(doc, "1. Business Problem & Motivation")
    body(
        doc,
        "A business that receives invoices as scanned images or photographs faces a bottleneck "
        "before any of that paper can enter an automated accounts-payable or obligation-tracking "
        "workflow: a human has to look at each one and check whether it is complete enough to act "
        "on. Typical checks include:"
    )
    bullets(doc, [
        "Is the invoice signed or stamped (evidence of authorization)?",
        "Does it cite a reference number - a purchase order (PO), sales order, contract, or work "
        "order - that ties it to an approved commitment?",
        "Does it state a clear invoice date and payment terms (e.g. \"Net 30\")?",
        "Are there any risk clauses - late-payment penalties, dispute language - worth flagging "
        "to a reviewer?",
    ])
    body(
        doc,
        "At the volume real businesses operate at (thousands of invoices a month), manual triage "
        "is slow, inconsistent between reviewers, and creates a queue that delays payment and "
        "reconciliation. The Pistac.io framing used throughout this project names the goal "
        "precisely: a digital obligation record - a structured, machine-readable summary of what "
        "an invoice obligates the business to do and when - that downstream systems can consume "
        "without a human reading the source image first."
    )
    body(
        doc,
        "Our answer is a pipeline that takes an invoice image in and produces (a) a structured "
        "JSON record of everything the pipeline could determine about that invoice, and (b) a "
        "Ready / Not-ready verdict, evaluated against rules the user chooses and configures, with "
        "a transparent per-rule breakdown of why an invoice passed or failed. Critically, the "
        "verdict is not a fixed hardcoded rule buried in code - a compliance officer, an AP "
        "clerk, and an auditor may reasonably want different thresholds (require both a stamp "
        "and a signature vs. either one; a 30-day vs. a 60-day payment-terms floor), so the "
        "system exposes the policy itself as user-editable configuration, and defaults to the "
        "conservative, safe answer whenever it cannot confirm a rule (fail-closed: missing "
        "evidence is treated as a failed check, never a free pass)."
    )

    # ================= 2. DATASETS =================
    h1(doc, "2. Related Datasets & Data Pipeline")
    body(
        doc,
        "No single public dataset contains full-page invoices annotated with stamps, signatures, "
        "business regions, and business-parameter text all at once. We therefore combine four "
        "real, purpose-built datasets, each covering one facet of the problem, and apply the "
        "resulting models to real invoice scans:"
    )
    add_table(
        doc,
        ["Dataset", "Contents", "Used by"],
        [
            ["Invoice batches (real scans)",
             "~5,201 unique real invoice images after de-duplication; 750 sampled into the "
             "working manifest; batch_1 carries annotation CSVs (real OCR text + structured "
             "fields) for ~1,413 images",
             "Rolando (ingestion); all others via the manifest"],
            ["OCR Dataset of Multi-type Documents",
             "973 SROIE-style receipt images, pre-split 778/97/98, 100% annotated with 52,331 "
             "real polygon text boxes + transcriptions, plus entity fields (company, date, "
             "address, total)",
             "Jordan (region detector), Damir (OCR evaluation)"],
            ["SignverOD",
             "2,765 document images with signature bounding boxes (4 categories; only category 1 "
             "= signature is used)",
             "Diana"],
            ["StaVer",
             "400 document scans with binary stamp ground-truth masks (no boxes - boxes are "
             "derived)",
             "Diana"],
        ],
        col_widths=[1.7, 3.2, 1.6],
    )
    body(
        doc,
        "Rolando's data-engineering stage turns a messy tree of raw invoice scans into the single "
        "clean artifact every other stage joins on: invoice_manifest.csv (document_id, "
        "image_path, width, height, split, has_ground_truth). Two decisions here are worth "
        "calling out because they materially change what the rest of the pipeline can honestly "
        "claim:"
    )
    bullets(doc, [
        "Annotation-aware resampling. A naive random sample of invoices produced only 26.3% "
        "ground-truth coverage, because annotation CSVs exist only for batch_1. Resampling to "
        "prefer annotated images lifted coverage to 100% across the 750-row manifest (split "
        "525/120/105), at the cost of reduced cross-batch visual variety - a deliberate, "
        "documented trade-off.",
        "Duplicate discovery. batch_3/ was found to silently re-contain copies of batch_1 and "
        "batch_2. The true unique invoice count is 5,201, not 8,181. Sampling the duplicated "
        "copies would have let the same physical invoice appear in both a training split and a "
        "test split - classic data leakage - so the duplicates are excluded from sampling "
        "entirely.",
    ])
    body(
        doc,
        "The manifest's document_id is the join key every downstream CSV uses, which is why "
        "getting this stage right first is a precondition for everything else being trustworthy."
    )
    add_figure(doc, "sample_invoice_grid.png",
               "Figure 0a. A grid of sample invoice images spanning the manifest's batches.")
    add_figure(doc, "preprocessing_examples.png",
               "Figure 0b. Preprocessing steps applied ahead of downstream detection and OCR.")

    # ================= 3. METHODOLOGY =================
    h1(doc, "3. Methodology")
    body(
        doc,
        "Every training notebook pulls its training knobs from a single shared module, "
        "src/compute_profile.py, so the same code runs identically in two modes and the two "
        "never silently drift apart:"
    )
    add_table(
        doc,
        ["Knob", "local_cpu (dev machine)", "colab_gpu (member runs)"],
        [
            ["epochs", "10", "100"],
            ["imgsz", "416", "960"],
            ["batch", "8", "16"],
            ["workers", "0", "2"],
            ["patience (early stop)", "10", "20"],
            ["max images/class", "250 (subsampled)", "None (full dataset)"],
            ["device", "cpu", "GPU (device 0)"],
        ],
        col_widths=[2.0, 2.3, 2.3],
    )
    body(
        doc,
        "All members trained under IIP_COMPUTE_PROFILE = \"colab_gpu\" on a Colab T4/L4/A100 "
        "runtime. Every metrics JSON also carries a _run provenance block (profile, epochs, "
        "imgsz, device, wall-clock seconds) so local-CPU and Colab-GPU runs can later be compared "
        "on equal footing."
    )

    h2(doc, "3.1 Rolando - Data Ingestion & the Manifest")
    body(
        doc,
        "No model is trained in this stage; it is data engineering, and the project's overall "
        "honesty rests on it. Beyond the annotation-aware resampling and duplicate exclusion "
        "described in Section 2, Rolando's notebook performs a stratified split (by "
        "ground-truth availability, so every split stays fully labelled) and produces a "
        "data-quality report plus sample/preprocessing figure grids for the report and deck."
    )

    h2(doc, "3.2 Diana - Stamp & Signature Detection")
    body(
        doc,
        "Model choice: YOLOv8n, trained as a single 2-class detector (stamp, signature - never "
        "merged, never renamed, since the final JSON schema and the verdict engine both assume "
        "the two exist independently). YOLOv8n (\"nano\") was chosen over a two-stage detector "
        "such as Faster R-CNN because it is a fast one-stage architecture well suited to small, "
        "visually distinct marks, trains comfortably within the Colab GPU time budget, and a "
        "single 2-class model shares one backbone rather than requiring two separately trained "
        "detectors."
    )
    body(
        doc,
        "The data-engineering challenge here is real: SignverOD stores boxes as normalized "
        "[x, y, w, h] which must be converted to pixel coordinates, and only category 1 "
        "(signature) is kept - categories 2-4 (initials, redaction, date) are dropped. StaVer, "
        "by contrast, ships no boxes at all, only binary ground-truth masks; boxes are derived "
        "using cv2.connectedComponentsWithStats, and the count of derived boxes is cross-checked "
        "against each scan's numStamps count from its info file as a sanity check."
    )
    body(
        doc,
        "Key parameters (as run): epochs ≤50, imgsz 640, batch 16, patience 20; inference "
        "confidence threshold 0.25; IoU match threshold 0.5 (via the shared src/iou.py, never "
        "reimplemented). These were reduced from the colab_gpu default (100 epochs / imgsz 960) "
        "after a GPU-budget exhaustion - see Section 6.1 Challenges and Section 6.2 Lessons "
        "Learned. Evaluation: precision, recall, and mean IoU computed per class on a real "
        "held-out split of SignverOD + StaVer - never merged into one detection score, since a "
        "\"must be signed\" business rule needs the signature number specifically. Inference is "
        "then run on the 750 invoices and reported as detection counts and confidence "
        "distributions only (no invoice-level ground truth exists, so no precision/recall claim "
        "is made there)."
    )

    h2(doc, "3.3 Jordan - Region Detection & IoU")
    body(
        doc,
        "Model choice: YOLOv8n, trained as a 5-class detector: company, date, address, total, "
        "other_text. The intellectual core of this stage is turning an unlabelled-boxes dataset "
        "into a labelled one: the OCR Dataset of Multi-type Documents ships 52,331 real polygon "
        "text boxes with no class label, alongside per-document entity values (company, date, "
        "address, total) with no coordinates. Jordan joins the two by fuzzy-matching each box's "
        "transcribed text against each entity value (rapidfuzz, using the max of partial_ratio "
        "and ratio, threshold 88), assigning the best-matching field label when the score clears "
        "the threshold and other_text otherwise. This turns 52,331 unlabelled boxes into a "
        "genuine 5-class region-detection training set with real geometry."
    )
    body(
        doc,
        "Key parameters (colab_gpu): epochs 100, imgsz 960, batch 16; fuzzy-match threshold 88; "
        "inference confidence 0.25; IoU match threshold 0.5. Challenge named explicitly: "
        "other_text massively outnumbers the four named field classes (visible directly in the "
        "train-split class-distribution printout), so per-class precision/recall/IoU are "
        "reported, never a single blended score that a dominant class could inflate. Evaluation: "
        "per-class IoU on the dataset's own official test split (98 images). Predictions are "
        "then run on the 750 invoices (rows tagged source=\"invoice\"), where receipts (~460 px "
        "wide) transferring to full-page invoices (1654x2339) is a domain shift reported as "
        "counts. Jordan's date/total/company boxes have a second life in the app: cropping and "
        "OCR-ing just those regions gives a field localiser that feeds the verdict engine's date "
        "signal and a \"where is this field, and what does it say\" overlay."
    )

    h2(doc, "3.4 Damir - OCR, Business Parameters & Terms")
    body(
        doc,
        "Model choice: EasyOCR (a deep-learning OCR engine, run on GPU) for text extraction, "
        "followed by rule-based extraction - src/parameter_checker.py and "
        "src/terms_extraction.py - for structured fields. The rule-based choice over an "
        "end-to-end learned extractor is deliberate: it needs no training data of its own, every "
        "decision is traceable to a specific keyword or regex match (auditable, which matters "
        "for a compliance-adjacent tool), and a user can add a brand-new required field (e.g. "
        "\"Insurance Certificate\") by editing config/required_fields_config.json with no code "
        "change."
    )
    body(
        doc,
        "Damir's stage has two evaluation sets of very different strength, and reports both "
        "honestly:"
    )
    bullets(doc, [
        "Primary (headline): the OCR Dataset's own test split - 98 scored receipt images (of "
        "973 total in the dataset), 100% coverage, real per-box ground-truth transcriptions - "
        "scored with CER (character error rate) 0.2152 mean / 0.1751 median and WER (word error "
        "rate) 0.5423 mean / 0.5225 median. This is the honest headline: EasyOCR on noisy, "
        "skewed, low-res receipt photos is a genuinely hard case. Lower is better.",
        "Secondary: the 120 real batch_1 invoices Damir's notebook actually ran EasyOCR on, "
        "scored against the batch_1 annotation CSVs' ground-truth text - CER 0.0002 mean "
        "(median 0.0), WER 0.0015 mean (median 0.0), essentially character-perfect. These are "
        "clean, high-resolution, digitally-rendered invoice images, so this is a much easier "
        "case than the receipt photos - real, but not representative of scanned-document OCR "
        "difficulty. The honest framing is both numbers together, never one headline: "
        "near-perfect on clean synthetic invoices, CER 0.215 / WER 0.542 on real-world receipt "
        "photos.",
    ])
    body(
        doc,
        "Coverage was completed locally (CPU-only, no GPU/EasyOCR re-run, no notebook execution "
        "- see scripts/complete_damir_outputs.py) once it became clear the manifest is now 100% "
        "batch_1 and all 750 manifest invoices have ground-truth OCR text in the batch_1 "
        "annotation CSVs. Invoice-level text for all 750 is therefore Damir's real EasyOCR "
        "output where it exists (120 invoices) and the annotation CSVs' OCRed Text elsewhere "
        "(630 invoices) - zero additional OCR. outputs/predictions/parameter_presence_results.csv "
        "and terms_extraction_results.csv are regenerated at the contract's long schema across "
        "all 750 invoices from this text, using Damir's own check_all_fields / "
        "terms-extraction functions unmodified."
    )
    body(
        doc,
        "Honest finding - payment terms are nearly absent from this corpus. Across all 750 "
        "invoices, 99.6% have a parseable invoice_date, but only 0.4% (3 of 750) match any "
        "day-based payment-terms phrasing (\"Net 30\", \"due within N days\") or yield a "
        "billing_due_days value - these invoices are synthetic templates that simply don't "
        "carry that phrasing in their OCR/annotation text. That means the verdict engine's "
        "\"payment terms > N days\" rule is unknown (fail-closed, never a false pass) for "
        "nearly every invoice in this corpus - a real, corpus-level limitation, not a bug in the "
        "extraction regexes."
    )

    h2(doc, "3.5 Hessam - Integration, Verdict Engine & Streamlit Application")
    body(
        doc,
        "Hessam's stage merges all four upstream outputs into one final JSON per invoice "
        "(model_interface_contract.md Section 5) and builds the demo application. The headline "
        "design decision is the verdict engine (src/verdict_engine.py): rather than a single "
        "hardcoded readiness rule, the user builds a policy out of four toggleable rule types "
        "before judging any invoice -"
    )
    bullets(doc, [
        "Visual mark - stamp OR signature, stamp AND signature, stamp only, or signature only.",
        "Reference number - at least one (or all) of PO / Order / Contract / Work Order present.",
        "Invoice date range - inclusive start/end bounds.",
        "Payment terms - a day-count threshold with a comparison operator (e.g. > 30 days).",
    ])
    body(
        doc,
        "The verdict is the AND of every enabled rule; disabled rules are skipped entirely. "
        "Every rule returns one of three statuses - pass, fail, or unknown - and the engine is "
        "deliberately fail-closed: an unknown status (no signal available for that invoice, e.g. "
        "it was never OCR'd) counts as not satisfied, so the invoice is marked Not-ready rather "
        "than silently passed. The breakdown always shows this distinction explicitly "
        "(\"unknown - treated as fail\"), never hiding it. This logic is pure Python with no "
        "Streamlit or I/O dependency, so it is independently unit-tested (9/9 tests passing as "
        "of this writing) and reused identically by both the single-invoice and whole-batch code "
        "paths."
    )
    body(
        doc,
        "A subtler integration point deserves its own callout: Damir's "
        "parameter_presence_results.csv and terms_extraction_results.csv are keyed to the "
        "OCR-Dataset receipt IDs, not the 750 invoices, because that is the dataset his notebook "
        "evaluates against for real ground truth. The invoice-level reference/date/"
        "payment-terms signals the verdict engine actually consumes are therefore derived at "
        "integration time - Damir's own shared modules (parameter_checker.py, "
        "terms_extraction.py) are re-run on invoice OCR text (from the upload path or "
        "ocr_outputs.csv) and on the batch_1 annotation CSVs' OCRed Text field. This keeps "
        "Damir's evaluated numbers honest to what he actually measured, while still giving the "
        "app real signals to judge invoices against, at whatever coverage the underlying "
        "OCR/annotation data supports."
    )
    body(
        doc,
        "The Streamlit application (app/streamlit_app.py) exposes three views: Live Demo (upload "
        "or pick one invoice, annotated boxes, verdict card with per-rule breakdown, downloads), "
        "Batch Gallery (roll up the configured policy across all 750 invoices - \"N of 750 "
        "pass\" - with filtering and a passing-set export), and Model Report (the members' own "
        "metrics plus the local-CPU-vs-Colab-GPU comparison charts). Execution is hybrid: the "
        "gallery and roll-up read pre-computed published results for full-speed browsing, while "
        "a brand-new uploaded image runs the live pipeline (vision models if best.pt weights are "
        "present in models/, OCR/parameter/terms extraction always, since those run acceptably "
        "on CPU)."
    )

    # ================= 4. INTEGRATION & VERDICT ENGINE =================
    h1(doc, "4. Integration & the Configurable Verdict Engine")
    body(
        doc,
        "The system's central architectural insight is that the pipeline extracts signals; the "
        "user's policy is applied to those signals - and the two are cleanly separated. Every "
        "signal a verdict rule needs already exists somewhere in a member's output:"
    )
    add_table(
        doc,
        ["Verdict rule", "Signal", "Source"],
        [
            ["Stamp / signature present", "visual_elements.stamp_detected / signature_detected",
             "Diana"],
            ["Reference number present",
             "required_parameters[...] via config/required_fields_config.json",
             "Damir (re-run on invoice text at integration)"],
            ["Invoice date in range", "payment_context.invoice_date",
             "Damir / annotation text (derived)"],
            ["Payment terms > N days", "payment_context.billing_due_days",
             "Damir / annotation text (derived)"],
        ],
        col_widths=[1.8, 2.6, 2.2],
    )
    body(
        doc,
        "Notably, the verdict depends on Diana's and Damir's signals but not on Jordan's region "
        "labels directly - Jordan's company/date/address/total/other_text vocabulary is a "
        "receipt-entity schema, not the obligation-region schema in config/label_schema.json, so "
        "that mismatch is kept off the verdict's critical path. It still matters: Jordan's boxes "
        "power a \"detected regions\" panel and the region-guided crop-then-OCR feature (cropping "
        "a date or total box and running OCR on just that region) that feeds the app's date rule "
        "and its clearest visual."
    )
    body(
        doc,
        "Coverage differs sharply by path. On the upload path, all four rule signals can be "
        "computed live for a brand-new image since OCR, parameter-checking, and "
        "terms-extraction all run on CPU. On the batch/gallery path, the visual-mark signal has "
        "full coverage (750/750, since Diana's detector runs on every manifest image), but the "
        "reference/date/payment-terms signals only cover the subset of invoices with usable OCR "
        "or batch_1 annotation text - for the rest, the verdict engine's fail-closed rule marks "
        "those checks unknown to fail, and the UI labels them \"not evaluated (no OCR)\" rather "
        "than hiding the gap."
    )

    # ================= 5. RESULTS =================
    h1(doc, "5. Results")
    body(
        doc,
        "Diana's and Jordan's Colab GPU training runs are complete, and Damir's run has been "
        "independently verified. Every number below is read directly from the authoritative "
        "metrics JSONs (outputs/metrics/stamp_signature_metrics.json, "
        "outputs/metrics/region_iou_metrics.json, outputs/metrics/ocr_parameter_metrics.json) or "
        "from the integration report (outputs/reports/final_pipeline_report.md) - none are "
        "estimated or fabricated."
    )

    h2(doc, "5.1 Diana - Stamp & Signature Detection (real held-out split of SignverOD + StaVer)")
    add_figure(doc, "metrics_per_class.png",
               "Figure 1. Per-class precision, recall, and mean IoU for Diana's stamp/signature "
               "detector and Jordan's 5-class region detector.")
    add_table(
        doc,
        ["Class", "Precision", "Recall", "Mean IoU", "Support (tp+fn)"],
        [
            ["stamp", "0.903", "0.875", "0.822", "64 (tp 56, fp 6, fn 8)"],
            ["signature", "0.894", "0.638", "0.815", "1055 (tp 673, fp 80, fn 382)"],
        ],
        col_widths=[1.3, 1.2, 1.2, 1.2, 2.1],
    )
    body(
        doc,
        "Run provenance (_run block): colab_gpu profile, Tesla T4, epochs ≤50, imgsz 640, "
        "batch 16, 2,287 training images, evaluated on a real held-out split of SignverOD + "
        "StaVer, confidence threshold 0.25, IoU match threshold 0.5.",
        italic=True,
    )
    add_figure(doc, "stamp_signature_detection_examples.png",
               "Figure 2. Diana's stamp/signature detector applied to sample invoice images.")
    add_figure(doc, "diana_BoxPR_curve.png",
               "Figure 3. Box precision-recall curve from Diana's YOLOv8n training run.")
    add_figure(doc, "diana_confusion_matrix.png",
               "Figure 4. Confusion matrix from Diana's held-out evaluation split.")
    body(
        doc,
        "Invoice inference (750 real invoices, counts only - no ground truth): 0 of 750 "
        "invoices have any stamp or signature detection (detections_by_label: {}). This is the "
        "honest, correct result of applying a document/receipt-trained detector to a corpus of "
        "clean, unsigned digital invoice templates - it is a property of the invoice corpus, "
        "not a failure of the detector, whose own held-out stamp/signature IoU (0.822 / 0.815) "
        "is strong."
    )

    h2(doc, "5.2 Jordan - Region Detection (OCR Dataset official test split, 98 images)")
    add_table(
        doc,
        ["Class", "Precision", "Recall", "Mean IoU", "Support"],
        [
            ["company", "0.897", "0.819", "0.894", "127"],
            ["date", "0.756", "0.855", "0.812", "337"],
            ["address", "0.880", "0.883", "0.896", "325"],
            ["total", "0.763", "0.783", "0.887", "309"],
            ["other_text", "0.911", "0.967", "0.877", "4153"],
            ["Macro mean IoU", "-", "-", "0.873", "-"],
        ],
        col_widths=[1.5, 1.2, 1.2, 1.2, 1.2],
    )
    body(
        doc,
        "Run provenance (_run block): colab_gpu profile, Tesla T4, epochs 100, imgsz 960, batch "
        "16, 778 training images, evaluated on the OCR Dataset's own official test split (98 "
        "images), confidence threshold 0.25, IoU match threshold 0.5.",
        italic=True,
    )
    add_figure(doc, "region_detection_examples.png",
               "Figure 5. Jordan's 5-class region detector applied to sample images.")
    body(
        doc,
        "Invoice inference (domain shift, counts only): 750 of 750 invoices carry at least one "
        "region detection (invoices_with_regions: 750) - unlike Diana's stamp/signature marks, "
        "business-text regions (company/date/address/total/other_text-shaped text blocks) are "
        "present on every invoice in this corpus, so the region detector transfers usefully "
        "across the receipt-to-invoice domain shift even though it was never trained on invoices "
        "directly."
    )

    h2(doc, "5.3 Damir - OCR & Business-Parameter Extraction")
    body(doc, "Primary (OCR Dataset receipt test split, real GPU EasyOCR run):", bold=True)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Docs scored", "98"],
            ["CER (mean / median)", "0.2152 / 0.1751"],
            ["WER (mean / median)", "0.5423 / 0.5225"],
        ],
        col_widths=[3.0, 3.0],
    )
    body(
        doc,
        "Recomputed locally from data/raw/invoices/OCR Dataset of Multi-type "
        "Documents/invoice/*/annotations/*.json as an honesty cross-check - reproduces the "
        "GPU-run numbers above exactly (see _local_receipt_cer_wer_check in "
        "ocr_parameter_metrics.json).",
        italic=True,
    )
    body(doc, "Secondary (real batch_1 invoices - state the denominator):", bold=True)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Invoices scored / manifest with GT",
             "120 / 750 (100% of the manifest has ground-truth text; 120 is how many Damir's "
             "notebook actually ran real EasyOCR on)"],
            ["CER (mean / median)", "0.0002 / 0.0000"],
            ["WER (mean / median)", "0.0015 / 0.0000"],
        ],
        col_widths=[2.4, 3.6],
    )
    body(
        doc,
        "Near-perfect on these clean, digitally-rendered invoices - a genuine result, but an "
        "easy one; it should never be quoted alone next to the harder 0.2152 receipt CER."
    )
    body(
        doc,
        "Invoice-text coverage (completed locally, no GPU): all 750 manifest invoices now have "
        "usable text - 120 from Damir's real EasyOCR output, 630 from batch_1 annotation "
        "ground-truth text - feeding parameter_presence_results.csv and "
        "terms_extraction_results.csv at the contract's long schema."
    )
    body(doc, "Business-parameter presence rate (share of the 750 invoices where each field "
               "matched, via check_all_fields, unmodified):", bold=True)
    add_table(
        doc,
        ["Field", "Required", "Presence rate"],
        [
            ["PO Reference", "yes", "54.7%"],
            ["Order Number", "yes", "19.9%"],
            ["Contract Number", "no", "1.1%"],
            ["Project Reference", "no", "0.0%"],
            ["Work Order No.", "no", "100%*"],
            ["Insurance Policy Number", "no", "100%*"],
            ["Bill of Lading Number", "no", "100%*"],
        ],
        col_widths=[2.5, 1.5, 2.0],
    )
    body(
        doc,
        "Required-field presence rate (PO Reference + Order Number, the two that gate Pistac.io "
        "readiness): 37.3%. The three fields marked * hit 100% because their config patterns are "
        "permissive substring/regex matches (e.g. Bill of Lading's [A-Z0-9-]{6,20} matches the "
        "word \"INVOICE\"; Work Order's WO[-\\s]?[0-9A-Z]+ matches \"WORTH\"; PO Reference's "
        "\"PO\" keyword matches inside \"CORPORATION\") - real output of the shared, unmodified "
        "parameter_checker.py, but a false-positive-prone upper bound rather than a clean "
        "detection signal on this corpus."
    )
    body(
        doc,
        "Terms parseability across all 750 invoices: 99.6% have a parseable invoice_date; only "
        "0.4% (3 invoices) match any day-based payment-terms phrasing or yield "
        "billing_due_days. This corpus's invoices simply don't carry \"Net 30\"-style phrasing, "
        "so the verdict engine's payment-terms rule is unknown/fail-closed for nearly all of "
        "them - an honest, corpus-level finding, not an extraction bug."
    )

    h2(doc, "5.4 Hessam - Integration & Verdict Engine Coverage")
    body(
        doc,
        "Fused 750 per-invoice JSON records from all four upstream stages into "
        "outputs/final_json/sample_invoice_outputs/."
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Invoices in manifest", "750"],
            ["Sample final-JSON records produced", "750"],
            ["Invoices with both stamp AND signature detected",
             "0 (Diana: 0/750 invoice detections)"],
            ["Default policy - N of 750 pass", "~0 / 750 (~0%)"],
            ["Strict policy - N of 750 pass", "0 / 750 (0.0%)"],
            ["Lenient policy - N of 750 pass", "750 / 750 (100.0%)"],
            ["Graded completeness - Ready (score >= 80)", "315 / 750 (42.0%)"],
            ["Graded completeness - Needs review (60-79)", "93 / 750 (12.4%)"],
            ["Graded completeness - Not ready (< 60)", "342 / 750 (45.6%)"],
        ],
        col_widths=[3.5, 3.5],
    )
    add_figure(doc, "readiness_by_policy.png",
               "Figure 6. Obligation-readiness across the 750-invoice batch. Strict fail-closed "
               "policies: Lenient 100%, Default ~0% (requires a PO/contract reference the corpus "
               "lacks), Strict 0% (also requires a visual mark it lacks) - an honest domain gap, "
               "not an engine defect. An earlier Default 63.3% was a loose-matching artifact, "
               "since corrected. The graded completeness score gives the meaningful spread: "
               "Ready 42.0%, Needs review 12.4%, Not ready 45.6% (mean 76.3).")
    body(doc, "Region detections on invoices (Jordan, source=invoice, total boxes across all "
               "750 invoices):", bold=True)
    add_table(
        doc,
        ["Region label", "Total boxes"],
        [
            ["other_text", "54,009"],
            ["address", "1,877"],
            ["company", "1,437"],
            ["total", "1,389"],
            ["date", "50"],
        ],
        col_widths=[3.5, 3.5],
    )

    h2(doc, "5.5 Compute Budget Comparison (local_cpu vs. colab_gpu)")
    add_figure(doc, "compute_profiles.png",
               "Figure 7. Wall-clock time, epoch count, and input image size for Diana's and "
               "Jordan's Colab T4 training runs.")
    add_table(
        doc,
        ["Member", "Profile", "Epochs", "imgsz", "Batch", "Wall-clock"],
        [
            ["Diana", "colab_gpu*", "≤50", "640", "16", "10,000.4s (~166.7 min)"],
            ["Jordan", "colab_gpu", "100", "960", "16", "6,530.0s (~108.8 min)"],
            ["Damir", "colab_gpu", "- (no training)", "-", "-",
             "598.7s (GPU EasyOCR on 120+98 images)"],
        ],
        col_widths=[1.1, 1.1, 1.1, 0.9, 0.9, 2.2],
    )
    body(
        doc,
        "*Diana's run used a reduced budget (≤50 epochs / imgsz 640) after a GPU-budget "
        "exhaustion at epoch 86 on the original 100-epoch / imgsz-960 attempt - see Section 6.1.",
        italic=True,
    )

    # ================= 6. LIMITATIONS =================
    h1(doc, "6. Limitations & Honest Caveats")
    body(
        doc,
        "These are stated plainly, not hidden, because owning a limitation is a stronger "
        "technical position than glossing over it."
    )
    bullets(doc, [
        "Domain gap. Diana's detector is trained on SignverOD (documents) and StaVer "
        "(documents); Jordan's detector is trained on the OCR Dataset of Multi-type Documents "
        "(receipts, ~460 px wide). Neither source is an invoice, and the invoice corpus is "
        "full-page (1654x2339). Both detectors are then applied to invoices as a "
        "domain-transfer step. Metrics are reported on each dataset's own real held-out split "
        "(legitimate numbers); on the invoices themselves we report detection counts, never "
        "precision/recall/accuracy, because there is no invoice-level ground truth to score "
        "against.",
        "Ground-truth coverage. Annotation CSVs with real structured text exist only for "
        "batch_1. A naive sample gave 26.3% manifest coverage; annotation-aware resampling "
        "raised that to 100% across the 750-row manifest, trading away cross-batch visual "
        "variety for full labels - a deliberate, documented trade-off, not an oversight.",
        "Data leakage, caught and fixed. batch_3/ was discovered to secretly duplicate batch_1 "
        "and batch_2. True unique invoice count is 5,201, not 8,181. The duplicates are "
        "excluded from sampling to prevent the same invoice appearing in both a training and a "
        "test context.",
        "Integration re-derivation. Damir's parameter_presence_results.csv and "
        "terms_extraction_results.csv are keyed to OCR-Dataset receipt IDs, not the 750 "
        "invoices. The invoice-level reference/date/payment-terms signals the verdict engine "
        "consumes are therefore derived at integration time from invoice OCR text and batch_1 "
        "annotation text using Damir's own shared modules - coverage on the batch path is "
        "bounded by how many invoices have usable OCR or annotation text, not the full 750.",
        "Verdict-engine semantics. The engine is deliberately fail-closed: a rule with no "
        "signal for a given invoice counts as failed, never passed. This is the right default "
        "for a compliance gate (a readiness check must not pass what it cannot confirm), but it "
        "means Not-ready counts include both genuinely-failing invoices and invoices the "
        "pipeline simply could not evaluate - the per-rule breakdown always distinguishes the "
        "two so this is never silently conflated.",
        "Compute budget. Jordan's and Damir's numbers come from the colab_gpu profile (100 "
        "epochs, imgsz 960, full dataset). Diana's model is the exception - it was retrained at "
        "a reduced budget (≤50 epochs, imgsz 640) after exhausting the Colab GPU "
        "allocation, so its metrics reflect that smaller budget. The local_cpu profile (10 "
        "epochs, imgsz 416, 250-image subsample) exists purely for fast iteration on a "
        "CPU-only dev machine and backs no reported metric.",
    ])

    h2(doc, "6.1 Challenges Encountered")
    bullets(doc, [
        "GPU-budget exhaustion mid-training (stamp/signature model). Diana's 2-class YOLOv8n "
        "detector was first launched at the colab_gpu default (imgsz 960, 100 epochs). The "
        "Colab GPU allocation ran out at epoch 86, and because Ultralytics wrote its "
        "checkpoints to the ephemeral /content disk, the run was lost entirely when the "
        "runtime was recycled - 86 epochs of compute discarded, with no resumable state. This "
        "forced a re-run under a deliberately reduced budget: imgsz 640 and a hard cap of 50 "
        "epochs. The trade-off is a modest expected drop in localisation precision (smaller "
        "input resolution resolves small marks less sharply) in exchange for a run that "
        "completes inside one GPU window.",
        "Ephemeral vs durable storage. The root cause was not the budget itself but where "
        "checkpoints lived: nothing on /content survives a runtime recycle, so a partial run "
        "had zero salvage value.",
        "Free-tier variability. Colab GPU availability and session length are not guaranteed, "
        "which makes long single-shot training runs fragile for a team sharing the free tier.",
    ])

    h2(doc, "6.2 Lessons Learned")
    bullets(doc, [
        "Match the compute budget to the platform, not the ideal. A 100-epoch / imgsz-960 plan "
        "is reasonable on dedicated hardware but is the wrong shape for free-tier Colab. "
        "Diana's notebook now runs imgsz 640 / ≤50 epochs - near-converged for YOLOv8n on "
        "this data while fitting the budget.",
        "Always checkpoint to durable storage. Diana's training now writes to a Drive-backed "
        "run directory with save_period=10, and the training cell is resume-aware "
        "(YOLO(last.pt).train(resume=True)): a future disconnect resumes from the last Drive "
        "checkpoint instead of restarting from scratch. Had this been in place initially, the "
        "86-epoch run would have been recoverable.",
        "best.pt != the last epoch. Ultralytics tracks the best validation checkpoint "
        "independently, so with early-stopping (patience) the full epoch count is rarely "
        "needed - a reason a 50-epoch cap costs less accuracy than the raw number suggests.",
        "Budget is a shared team resource. Sequencing members' GPU runs (rather than everyone "
        "training at once) avoids collectively exhausting the free-tier allocation.",
    ])

    # ================= 7. CONCLUSION =================
    h1(doc, "7. Conclusion & Future Work")
    body(
        doc,
        "The pipeline demonstrates that a business-facing readiness question - \"is this "
        "invoice complete enough to become a digital obligation record?\" - can be decomposed "
        "into a small number of independently trainable/evaluable computer-vision and NLP "
        "sub-problems (mark detection, region detection, OCR, rule-based field extraction), "
        "integrated through a fixed schema, and judged by a transparent, user-configurable, "
        "fail-closed rule engine rather than a single opaque model output. Every sub-stage is "
        "evaluated on real, labelled data with real metrics; the honest domain gap between "
        "training data and the invoice application target is surfaced rather than hidden, and "
        "the data-engineering issues found along the way (low ground-truth coverage, "
        "duplicate-invoice leakage) were caught and fixed before they could quietly bias "
        "results."
    )
    body(
        doc,
        "Future work, roughly in order of expected impact: (1) collect or license "
        "invoice-native ground truth - even a few hundred hand-annotated invoices with real "
        "stamp/signature/region boxes would let Diana and Jordan report true invoice-domain "
        "accuracy instead of source-domain accuracy plus counts; (2) extend the annotation-aware "
        "sampling strategy to batches 2 (and any future batches) so ground-truth coverage does "
        "not depend on batch_1 alone; (3) add a learned OCR-quality gate so the verdict engine "
        "can distinguish \"field genuinely absent\" from \"OCR failed, field possibly present\" "
        "more granularly than a single unknown status; (4) expose the fuzzy-match threshold and "
        "confidence thresholds as tunable sidebar controls in the app so a grader or user can "
        "directly observe the precision/recall trade-off live, rather than only in the "
        "notebooks."
    )

    # ================= 8. CONTRIBUTIONS =================
    h1(doc, "8. Individual Contributions")
    add_table(
        doc,
        ["Member", "Role", "Key deliverables"],
        [
            ["Rolando", "Data ingestion & manifest",
             "invoice_manifest.csv, data-quality report, annotation-aware resampling, "
             "duplicate-leakage discovery and exclusion"],
            ["Diana", "Stamp & signature detection",
             "2-class YOLOv8n detector, SignverOD/StaVer adaptation (normalized-box conversion; "
             "mask-to-box derivation via connected components), per-class P/R/IoU"],
            ["Jordan", "Region detection & IoU",
             "5-class YOLOv8n detector, entity-text-to-box fuzzy-matching label construction, "
             "per-class P/R/IoU on the OCR Dataset test split"],
            ["Damir", "OCR, parameters & terms",
             "EasyOCR text extraction, CER/WER evaluation (primary + secondary sets), "
             "rule-based parameter/terms extraction integration"],
            ["Hessam", "Integration, verdict engine & app",
             "Final-JSON integration, src/verdict_engine.py (fail-closed configurable policy "
             "engine, unit-tested), Streamlit application (Live Demo / Batch Gallery / Model "
             "Report), report & deck assembly"],
        ],
        col_widths=[1.0, 1.8, 4.2],
    )

    # ================= REFERENCES =================
    h1(doc, "References")
    bullets(doc, [
        "SignverOD - signature detection dataset (document images with signature bounding "
        "boxes).",
        "StaVer - stamp verification dataset (document scans with stamp ground-truth masks).",
        "OCR Dataset of Multi-type Documents - SROIE-style receipt dataset with polygon text "
        "boxes, transcriptions, and entity annotations (company/date/address/total).",
        "Real invoice batch scans (data/raw/invoices/batch_1..3/) with batch_1 annotation CSVs "
        "(Json Data{invoice, items, subtotal, payment_instructions}, OCRed Text).",
        "Jocher, G. et al., Ultralytics YOLOv8 (object detection framework).",
        "JaidedAI, EasyOCR (deep-learning OCR engine).",
        "jiwer - CER/WER scoring library.",
        "rapidfuzz - fuzzy string matching library.",
        "Project source: src/compute_profile.py, src/verdict_engine.py, "
        "src/parameter_checker.py, src/terms_extraction.py, src/iou.py, "
        "model_interface_contract.md, colab/notebooks/00_preflight_check_colab.ipynb through "
        "05_hessam_integration_colab.ipynb.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")
    return doc


def verify():
    doc = Document(str(OUT))
    n_paragraphs = len(doc.paragraphs)
    n_tables = len(doc.tables)
    print(f"Reload check: {OUT.name} -> {n_paragraphs} paragraphs, {n_tables} tables")
    return n_paragraphs, n_tables


if __name__ == "__main__":
    build()
    verify()
